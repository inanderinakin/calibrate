import boto3
from dotenv import load_dotenv
import re
import time
import docx

load_dotenv()

client = boto3.client("textract")

def extract_docx_text(file_path: str) -> list[str]:
    doc = docx.Document(file_path)
    lines = []
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in lines:
                    lines.append(text)
                    
    return lines

def extract_cv_text(bucket: str, key: str) -> list[str]:
    response = client.start_document_text_detection(
        DocumentLocation = {
            'S3Object': {
            'Bucket': bucket,
            'Name': key
            }
        }
    )

    isSuccess = ""
    attemptCount= 0
    print("Waiting for results...")
    while (isSuccess != "SUCCEEDED"):
        attemptCount += 1
        if (attemptCount< 60):
            result = client.get_document_text_detection(JobId=response["JobId"])
            isSuccess = result["JobStatus"]

            if isSuccess == "SUCCEEDED":
                print("Success!")
                break

            elif isSuccess == "FAILED":
                raise RuntimeError(f"Request failed: {result.get('StatusMessage')}") 

            time.sleep(5)
        else:
            raise RuntimeError("Max attempt reached. Please try again later...")

    sidebar_lines = [block["Text"] for block in result["Blocks"] if block["BlockType"] == "LINE" and block["Geometry"]["BoundingBox"]["Left"] < 0.2]
    main_lines = [block["Text"] for block in result["Blocks"] if block["BlockType"] == "LINE" and block["Geometry"]["BoundingBox"]["Left"] >= 0.2]
    return main_lines + sidebar_lines

def extract_skill_candidates(lines: list[str]) -> list[str]:
    potential_candidates = []

    for i in range (len(lines) - 1):
        zippedLine = lines[i] + " " + lines[i+1]
        potential_candidates.append(zippedLine)

    for line in lines:
        for fragment in re.split(r"[:,/()]", line):
            fragment = fragment.strip()
            if (len(fragment) > 2 and len(fragment) < 40):
                potential_candidates.append(fragment)
            
        for fragment in line.split():
            fragment = fragment.strip(".,;:!?()\"'")
            potential_candidates.append(fragment)
    
    return list(dict.fromkeys(potential_candidates))